from __future__ import annotations

import asyncio
import copy
from dataclasses import fields, replace
import hashlib
import json
import os
from pathlib import Path
import stat
from types import MappingProxyType
from typing import Any

from docx import Document
from PIL import Image
import pytest

from app import release_features
from app.office_validation import (
    EvidenceBox,
    OfficeDraftSeal,
    OfficeDraftValidationResult,
    OfficePrecommitRepairError,
    OfficePrecommitRepairRequest,
    OfficePrecommitRequest,
    OfficeValidationReport,
    ValidationCheck,
    build_precommit_repair_request,
    copy_replacement_args,
)
from app.schemas.agent import AgentInfo
from app.tool.builtin import office as office_module
from app.tool.builtin.office import OfficeTool
from app.tool.context import ToolContext
from app.tool.workspace_transaction import (
    WorkspaceMutationTransaction,
    WorkspaceOfficePrecommitView,
    WorkspacePrecommitSealError,
)


def _context(
    workspace: Path,
    coordinator: object,
    *,
    repairer: object | None = None,
) -> ToolContext:
    context = ToolContext(
        session_id="repair-session",
        message_id="repair-message",
        agent=AgentInfo(name="test", description="", mode="primary"),
        call_id="repair-call",
        language="en",
        workspace=str(workspace),
        root_turn_id="repair-root-turn",
        turn_run_id="repair-turn-run",
        checkpoint_id="repair-checkpoint",
        workspace_instance_id="repair-workspace-instance",
    )
    state: dict[str, object] = {"office_precommit_coordinator": coordinator}
    if repairer is not None:
        state["office_precommit_repairer"] = repairer
    context._app_state = state  # type: ignore[attr-defined]
    return context


@pytest.fixture
def workspace(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    private = tmp_path / "private"
    private.mkdir()
    monkeypatch.setenv("SUXIAOYOU_PRIVATE_DATA_DIR", str(private))
    return workspace


@pytest.fixture
def office_authoring_gate(monkeypatch: pytest.MonkeyPatch) -> None:
    for gate in (
        "V11_CHECKPOINTS_RELEASED",
        "V11_REWIND_RELEASED",
        "V11_VALIDATION_AGENT_RELEASED",
        "V11_OFFICE_V2_RELEASED",
    ):
        monkeypatch.setattr(release_features, gate, True)


def _candidate_result(
    request: OfficePrecommitRequest,
    view: WorkspaceOfficePrecommitView,
    outcome: str,
) -> OfficeDraftValidationResult:
    payload = view.staged_target.read_bytes()
    root_info = view.staged_root.lstat()
    source_info = view.staged_target.lstat()
    digest = hashlib.sha256(payload).hexdigest()
    baseline_sha256 = (
        view.baseline.sha256
        if view.baseline is not None and view.baseline.sha256 is not None
        else "0" * 64
    )
    check = ValidationCheck(
        code="layout_overflow",
        outcome=outcome,  # type: ignore[arg-type]
        message=f"private diagnostic containing {view.staged_target}",
        box=(
            EvidenceBox(page_number=2, x=10, y=20, width=30, height=40)
            if outcome != "pass"
            else None
        ),
        metrics=(("private_threshold", 0.5),),
    )
    report = OfficeValidationReport(
        document_format=request.document_format,
        baseline_sha256=baseline_sha256,
        candidate_sha256=digest,
        renderer_id="private-renderer",
        renderer_version="private-version",
        font_digest="f" * 64,
        verdict=outcome,  # type: ignore[arg-type]
        checks=(check,),
        checkpoint_id=request.checkpoint_id,
        root_turn_id=request.root_turn_id,
    )
    seal = OfficeDraftSeal(
        relative_path=view.relative_path,
        source_sha256=digest,
        source_mode=stat.S_IMODE(source_info.st_mode),
        source_size=source_info.st_size,
        root_identity=(root_info.st_dev, root_info.st_ino),
        source_identity=(source_info.st_dev, source_info.st_ino),
        validation_generation=view.validation_generation,
        renderer_id=report.renderer_id,
        renderer_version=report.renderer_version,
        font_digest=report.font_digest,
        parameters_version="private-parameters",
        parameters_sha256="a" * 64,
        quality="authoritative",
        cache_key=hashlib.sha256(b"private-cache" + payload).hexdigest(),
    )
    return OfficeDraftValidationResult(report=report, candidate=seal)


class _Session:
    def __init__(
        self,
        request: OfficePrecommitRequest,
        view: WorkspaceOfficePrecommitView,
        outcome: str,
    ) -> None:
        self.request = request
        self.view = view
        self.outcome = outcome
        self.state = "begun"
        self.result: OfficeDraftValidationResult | None = None

    async def validate_candidate(self) -> OfficeDraftValidationResult:
        assert self.state == "begun"
        self.result = _candidate_result(self.request, self.view, self.outcome)
        self.state = "validated"
        return self.result

    def consume_commit_seal(
        self,
        result: OfficeDraftValidationResult,
    ) -> OfficeDraftSeal:
        assert self.state == "validated"
        assert result is self.result
        seal = result.commit_seal
        assert seal is not None
        self.state = "committing"
        return seal

    def mark_committed(self, result: OfficeDraftValidationResult) -> None:
        assert self.state == "committing"
        assert result is self.result
        self.state = "committed"

    def abort(self) -> None:
        if self.state != "committed":
            self.state = "aborted"


class _Coordinator:
    def __init__(self, outcomes: list[str]) -> None:
        self.outcomes = list(outcomes)
        self.sessions: list[_Session] = []
        self.views: list[WorkspaceOfficePrecommitView] = []

    async def begin(
        self,
        *,
        request: OfficePrecommitRequest,
        view: WorkspaceOfficePrecommitView,
    ) -> _Session:
        outcome = self.outcomes[len(self.sessions)]
        session = _Session(request, view, outcome)
        self.sessions.append(session)
        self.views.append(view)
        return session


class _Repairer:
    def __init__(
        self,
        replacements: list[dict[str, Any]],
        *,
        before_return: Any | None = None,
        preserve_raw_target: bool = False,
    ) -> None:
        self.replacements = replacements
        self.before_return = before_return
        self.preserve_raw_target = preserve_raw_target
        self.requests: list[OfficePrecommitRepairRequest] = []

    def _replacement_for_request(
        self,
        request: OfficePrecommitRepairRequest,
        index: int,
    ) -> dict[str, Any]:
        replacement = copy.deepcopy(self.replacements[index])
        if not self.preserve_raw_target:
            replacement["file_path"] = request.tokenized_args["file_path"]
        return replacement

    async def repair(
        self,
        request: OfficePrecommitRepairRequest,
    ) -> dict[str, Any]:
        self.requests.append(request)
        if self.before_return is not None:
            self.before_return(len(self.requests), request)
        return self._replacement_for_request(request, len(self.requests) - 1)


class _BlockingRepairer(_Repairer):
    def __init__(self, replacement: dict[str, Any]) -> None:
        super().__init__([replacement])
        self.entered = asyncio.Event()
        self.release = asyncio.Event()

    async def repair(
        self,
        request: OfficePrecommitRepairRequest,
    ) -> dict[str, Any]:
        self.requests.append(request)
        self.entered.set()
        await self.release.wait()
        return self._replacement_for_request(request, 0)


class _CancellationSuppressingRepairer(_Repairer):
    """Never finishes until the test releases it after bounded detachment."""

    def __init__(self, replacement: dict[str, Any]) -> None:
        super().__init__([replacement])
        self.entered = asyncio.Event()
        self.cancelled = asyncio.Event()
        self.release = asyncio.Event()
        self.finished = asyncio.Event()

    async def repair(
        self,
        request: OfficePrecommitRepairRequest,
    ) -> dict[str, Any]:
        self.requests.append(request)
        self.entered.set()
        try:
            await asyncio.Event().wait()
        except asyncio.CancelledError:
            self.cancelled.set()
            await self.release.wait()
            self.finished.set()
            # The detached-task callback must retrieve this without logging its
            # deliberately private text.
            raise RuntimeError("private repair failure /staging/secret")


class _SlowRepairer(_Repairer):
    def __init__(self, replacement: dict[str, Any], delay: float) -> None:
        super().__init__([replacement])
        self.delay = delay

    async def repair(
        self,
        request: OfficePrecommitRepairRequest,
    ) -> dict[str, Any]:
        self.requests.append(request)
        await asyncio.sleep(self.delay)
        return self._replacement_for_request(request, 0)


def _install_test_repair_limits(
    monkeypatch: pytest.MonkeyPatch,
    *,
    timeout: float,
    grace: float,
) -> None:
    real_await = office_module._await_repair_worker

    async def bounded(repairer, request):
        return await real_await(
            repairer,
            request,
            timeout_seconds=timeout,
            settlement_grace_seconds=grace,
        )

    monkeypatch.setattr(office_module, "_await_repair_worker", bounded)


def _create_args(title: str) -> dict[str, Any]:
    return {
        "file_path": "report.docx",
        "operation": "create",
        "overwrite": False,
        "document": {"title": title},
    }


def _edit_args(old: str, new: str) -> dict[str, Any]:
    return {
        "file_path": "report.docx",
        "operation": "edit",
        "replacements": [{"old_text": old, "new_text": new}],
    }


def _create_layout_args(title: str, margin: float) -> dict[str, Any]:
    args = _create_args(title)
    args["document"]["sections"] = [
        {
            "action": "configure",
            "index": 0,
            "margins": {"left_inches": margin},
        }
    ]
    return args


def _edit_layout_args(old: str, new: str, margin: float) -> dict[str, Any]:
    args = _edit_args(old, new)
    args["document"] = {
        "sections": [
            {
                "action": "configure",
                "index": 0,
                "margins": {"left_inches": margin},
            }
        ]
    }
    return args


def _target(workspace: Path) -> Path:
    return workspace / "suxiaoyou_written" / "report.docx"


@pytest.mark.parametrize(
    ("original", "replacement"),
    (
        (
            _create_args("Board-approved revenue: 100"),
            _create_args("Changed revenue: 0"),
        ),
        (
            {
                "file_path": "book.xlsx",
                "operation": "create",
                "workbook": {
                    "cells": [
                        {"sheet": "Sheet1", "cell": "A1", "value": "=SUM(B1:B9)"}
                    ]
                },
            },
            {
                "file_path": "book.xlsx",
                "operation": "create",
                "workbook": {
                    "cells": [
                        {"sheet": "Sheet1", "cell": "A1", "value": "=0"}
                    ]
                },
            },
        ),
        (
            {
                "file_path": "deck.pptx",
                "operation": "create",
                "presentation": {
                    "slides": [{"title": "Status", "speaker_notes": "Risk is high"}]
                },
            },
            {
                "file_path": "deck.pptx",
                "operation": "create",
                "presentation": {
                    "slides": [{"title": "Status", "speaker_notes": "No risk"}]
                },
            },
        ),
        (
            {
                "file_path": "deck.pptx",
                "operation": "create",
                "first_party_template": {
                    "template_id": "status-update",
                    "template_version": "1.0.0",
                    "values": {"summary": "Board-approved revenue: 100"},
                },
            },
            {
                "file_path": "deck.pptx",
                "operation": "create",
                "first_party_template": {
                    "template_id": "status-update",
                    "template_version": "1.0.0",
                    "values": {"summary": "Changed revenue: 0"},
                },
            },
        ),
    ),
)
def test_automatic_repair_cannot_change_semantic_content(
    workspace: Path,
    original: dict[str, Any],
    replacement: dict[str, Any],
) -> None:
    with pytest.raises(
        OfficePrecommitRepairError,
        match="semantic content",
    ):
        office_module._validate_office_replacement_args(
            replacement,
            original=original,
            workspace=str(workspace),
            original_read_paths=frozenset(),
        )


def test_automatic_repair_can_change_only_reviewed_layout_fields(
    workspace: Path,
) -> None:
    original = _create_layout_args("fixed content", 1.5)
    replacement = _create_layout_args("fixed content", 1.0)

    assert office_module._validate_office_replacement_args(
        replacement,
        original=original,
        workspace=str(workspace),
        original_read_paths=frozenset(),
    ) is replacement


def _document_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


_REPAIR_RUNTIME_IDS = (
    "repair-session",
    "repair-message",
    "repair-call",
    "repair-root-turn",
    "repair-turn-run",
    "repair-checkpoint",
    "repair-workspace-instance",
)


def _privacy_source_report(*private_values: str) -> OfficeValidationReport:
    return OfficeValidationReport(
        document_format="docx",
        baseline_sha256="1" * 64,
        candidate_sha256="2" * 64,
        renderer_id="private-renderer",
        renderer_version="private-version",
        font_digest="3" * 64,
        verdict="fail",
        checks=(
            ValidationCheck(
                code="layout_overflow",
                outcome="fail",
                message=" | ".join((*_REPAIR_RUNTIME_IDS, *private_values)),
                metrics=(("private_threshold", 0.5),),
            ),
        ),
        checkpoint_id="repair-checkpoint",
        root_turn_id="repair-root-turn",
    )


def _repair_request_payload(request: OfficePrecommitRepairRequest) -> dict[str, Any]:
    return {
        "tokenized_args": copy_replacement_args(request.tokenized_args),
        "attempt": request.attempt,
        "report": {
            "document_format": request.report.document_format,
            "verdict": request.report.verdict,
            "checks": [
                {
                    "code": check.code,
                    "outcome": check.outcome,
                    "box": (
                        None
                        if check.box is None
                        else {
                            "page_number": check.box.page_number,
                            "x": check.box.x,
                            "y": check.box.y,
                            "width": check.box.width,
                            "height": check.box.height,
                        }
                    ),
                }
                for check in request.report.checks
            ],
        },
    }


def _path_rich_repair_args(workspace: Path) -> dict[str, Any]:
    relative_image = "素材 目录/图 像.png"
    absolute_image = str(workspace / "绝对 素材" / "另 一张.png")
    return {
        "file_path": str(workspace / "输出 目录" / "季度 报告.docx"),
        "operation": "create",
        "overwrite": False,
        "document": {
            "title": "safe title",
            "images": [
                {"path": relative_image, "caption": "one"},
                {"path": relative_image, "caption": "two"},
                {"path": absolute_image, "caption": "three"},
            ],
            "charts": [
                {
                    "path": relative_image,
                    "alt_text": "chart",
                    "source": "local",
                }
            ],
        },
        "presentation": {
            "slides": [
                {"images": [{"path": absolute_image}]},
            ],
        },
    }


def test_repair_path_tokens_are_domain_separated_shared_and_path_free(
    workspace: Path,
) -> None:
    original = _path_rich_repair_args(workspace)
    tokenized, token_table = office_module._tokenize_office_repair_args(original)
    request = build_precommit_repair_request(
        tokenized_args=tokenized,
        report=_privacy_source_report(
            original["file_path"],
            original["document"]["images"][0]["path"],
            original["document"]["images"][2]["path"],
        ),
        attempt=1,
    )

    prefix = office_module._OFFICE_REPAIR_TOKEN_PREFIX
    target_token = request.tokenized_args["file_path"]
    first_read = request.tokenized_args["document"]["images"][0]["path"]
    duplicate_read = request.tokenized_args["document"]["images"][1]["path"]
    absolute_read = request.tokenized_args["document"]["images"][2]["path"]
    chart_read = request.tokenized_args["document"]["charts"][0]["path"]
    presentation_read = request.tokenized_args["presentation"]["slides"][0][
        "images"
    ][0]["path"]
    assert isinstance(target_token, str) and target_token.startswith(prefix + "target:")
    assert isinstance(first_read, str) and first_read.startswith(prefix + "read:")
    assert first_read == duplicate_read == chart_read
    assert presentation_read == absolute_read
    assert absolute_read != first_read
    assert target_token not in {first_read, absolute_read}
    target_nonce = target_token.removeprefix(prefix + "target:")
    read_nonce, read_index = first_read.removeprefix(prefix + "read:").rsplit(":", 1)
    absolute_nonce, absolute_index = absolute_read.removeprefix(
        prefix + "read:"
    ).rsplit(":", 1)
    assert target_nonce == read_nonce == absolute_nonce
    assert {read_index, absolute_index} == {"0", "1"}

    second_tokenized, _second_table = office_module._tokenize_office_repair_args(
        original
    )
    assert second_tokenized["file_path"] != target_token

    restored = office_module._unmask_office_replacement_args(
        copy_replacement_args(request.tokenized_args),
        token_table,
    )
    assert restored == original

    serialized = json.dumps(
        _repair_request_payload(request),
        ensure_ascii=False,
        sort_keys=True,
    )
    path_values = (
        original["file_path"],
        original["document"]["images"][0]["path"],
        original["document"]["images"][2]["path"],
        str(workspace),
    )
    assert all(value not in serialized for value in path_values)
    assert all(runtime_id not in serialized for runtime_id in _REPAIR_RUNTIME_IDS)
    assert "private-renderer" not in serialized
    assert "private_threshold" not in serialized


@pytest.mark.parametrize(
    "attack",
    ["unknown", "modified", "replay", "swap", "new-read"],
)
def test_repair_path_tokens_reject_forgery_replay_swap_and_new_reads(
    workspace: Path,
    attack: str,
) -> None:
    original = _path_rich_repair_args(workspace)
    tokenized, token_table = office_module._tokenize_office_repair_args(original)
    replacement = copy_replacement_args(tokenized)
    selected_table = token_table
    if attack == "unknown":
        replacement["document"]["images"][0]["path"] = (
            office_module._OFFICE_REPAIR_TOKEN_PREFIX + "read:not-issued:0"
        )
    elif attack == "modified":
        replacement["file_path"] += "-modified"
    elif attack == "replay":
        _other_args, selected_table = office_module._tokenize_office_repair_args(
            original
        )
    elif attack == "swap":
        replacement["file_path"] = replacement["document"]["images"][0][
            "path"
        ]
    else:
        replacement["document"]["images"].append(
            {"path": str(workspace / "新 读取.png")}
        )

    with pytest.raises(OfficePrecommitRepairError):
        office_module._unmask_office_replacement_args(
            replacement,
            selected_table,
        )


@pytest.mark.parametrize(
    "structured_payload",
    [
        {"document": {"sections": [{"background_path": "private.png"}]}},
        {"workbook": {"sheets": [{"source_paths": ["private.csv"]}]}},
        {
            "presentation": {
                "slides": [{"speaker_notes_path": "private.txt"}]
            }
        },
    ],
)
def test_repair_tokenization_rejects_unsupported_structured_path_fields(
    structured_payload: dict[str, Any],
) -> None:
    args = {
        "file_path": "report.docx",
        "operation": "create",
        **structured_payload,
    }

    with pytest.raises(
        OfficePrecommitRepairError,
        match="unsupported path field",
    ):
        office_module._tokenize_office_repair_args(args)


@pytest.mark.asyncio
async def test_create_repairs_once_in_private_staging(
    workspace: Path,
    office_authoring_gate: None,
) -> None:
    original = _create_layout_args("fixed content", 1.5)
    replacement = _create_layout_args("fixed content", 1.0)
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _Repairer([replacement])

    result = await OfficeTool().execute(
        original,
        _context(workspace, coordinator, repairer=repairer),
    )

    assert result.success
    assert "fixed content" in _document_text(_target(workspace))
    assert len(coordinator.sessions) == 2
    assert [item.state for item in coordinator.sessions] == ["aborted", "committed"]
    assert coordinator.views[0] is not coordinator.views[1]
    assert coordinator.views[0].staged_root == coordinator.views[1].staged_root
    assert (
        coordinator.views[0].validation_generation
        != coordinator.views[1].validation_generation
    )
    assert result.metadata["office_validation_repair_attempts"] == 1
    request = repairer.requests[0]
    assert request.attempt == 1
    assert isinstance(request.tokenized_args, MappingProxyType)
    assert set(field.name for field in fields(request)) == {
        "tokenized_args",
        "report",
        "attempt",
    }
    assert set(field.name for field in fields(request.report)) == {
        "document_format",
        "verdict",
        "checks",
    }
    assert request.report.checks[0].code == "layout_overflow"
    assert request.report.checks[0].box is not None
    with pytest.raises(TypeError):
        request.tokenized_args["operation"] = "edit"  # type: ignore[index]
    with pytest.raises(TypeError):
        request.tokenized_args["document"]["title"] = "mutated"  # type: ignore[index]


@pytest.mark.asyncio
async def test_edit_reset_replays_repair_from_visible_baseline(
    workspace: Path,
    office_authoring_gate: None,
) -> None:
    target = _target(workspace)
    target.parent.mkdir()
    document = Document()
    office_module._drop_default_docx_custom_xml(document)
    document.add_paragraph("baseline")
    document.save(target)
    original = _edit_layout_args("baseline", "updated", 1.5)
    replacement = _edit_layout_args("baseline", "updated", 1.0)
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _Repairer([replacement])

    result = await OfficeTool().execute(
        original,
        _context(workspace, coordinator, repairer=repairer),
    )

    assert result.success
    assert "updated" in _document_text(target)
    assert len(coordinator.sessions) == 2
    assert coordinator.views[0].staged_root_identity == (
        coordinator.views[1].staged_root_identity
    )
    assert (
        coordinator.views[0].validation_generation
        != coordinator.views[1].validation_generation
    )


@pytest.mark.asyncio
async def test_two_repairs_can_pass_but_a_third_is_never_called(
    workspace: Path,
    office_authoring_gate: None,
) -> None:
    original = _create_layout_args("fixed content", 1.5)
    repairer = _Repairer(
        [
            _create_layout_args("fixed content", 1.25),
            _create_layout_args("fixed content", 1.0),
        ]
    )
    coordinator = _Coordinator(["fail", "fail", "pass"])

    result = await OfficeTool().execute(
        original,
        _context(workspace, coordinator, repairer=repairer),
    )

    assert result.success
    assert "fixed content" in _document_text(_target(workspace))
    assert [request.attempt for request in repairer.requests] == [1, 2]
    assert (
        copy_replacement_args(repairer.requests[0].tokenized_args)
        == copy_replacement_args(repairer.requests[1].tokenized_args)
    )
    assert len(coordinator.sessions) == 3
    assert len({id(item) for item in coordinator.sessions}) == 3


@pytest.mark.asyncio
async def test_third_failed_candidate_returns_redacted_round_metadata(
    workspace: Path,
    office_authoring_gate: None,
) -> None:
    repairer = _Repairer(
        [
            _create_layout_args("fixed content", 1.25),
            _create_layout_args("fixed content", 1.0),
        ]
    )
    coordinator = _Coordinator(["fail", "fail", "fail"])

    result = await OfficeTool().execute(
        _create_layout_args("fixed content", 1.5),
        _context(workspace, coordinator, repairer=repairer),
    )

    assert not result.success
    assert not _target(workspace).exists()
    assert len(repairer.requests) == 2
    assert len(coordinator.sessions) == 3
    failure = result.metadata["office_validation_failure"]
    assert failure["rounds"] == 3
    assert failure["repair_attempts"] == 2
    assert failure["repair_status"] == "limit_reached"
    assert failure["checks"] == [
        {
            "round": round_number,
            "code": "layout_overflow",
            "outcome": "fail",
            "page_number": 2,
            "box": {"x": 10, "y": 20, "width": 30, "height": 40},
        }
        for round_number in (1, 2, 3)
    ]
    serialized = json.dumps(result.metadata, sort_keys=True)
    assert str(workspace) not in serialized
    for forbidden in (
        "cache_key",
        "candidate_sha256",
        "golden",
        "private-renderer",
        "private-threshold",
        "seal",
        "staging",
    ):
        assert forbidden not in serialized


@pytest.mark.asyncio
async def test_missing_repairer_fails_after_one_candidate(
    workspace: Path,
    office_authoring_gate: None,
) -> None:
    coordinator = _Coordinator(["fail", "pass"])

    result = await OfficeTool().execute(
        _create_args("bad"),
        _context(workspace, coordinator),
    )

    assert not result.success
    assert len(coordinator.sessions) == 1
    assert result.metadata["office_validation_failure"]["repair_status"] == (
        "unavailable"
    )
    assert not _target(workspace).exists()


@pytest.mark.asyncio
async def test_repair_exception_log_exposes_only_exception_type(
    workspace: Path,
    office_authoring_gate: None,
    caplog: pytest.LogCaptureFixture,
) -> None:
    private_values = (
        str(workspace / "私密 目录" / "报告.docx"),
        *_REPAIR_RUNTIME_IDS,
        "private-exception-marker",
    )

    def explode(_attempt: int, _request: OfficePrecommitRepairRequest) -> None:
        raise RuntimeError(" | ".join(private_values))

    caplog.set_level("WARNING", logger=office_module.__name__)
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _Repairer(
        [_create_args("unused")],
        before_return=explode,
    )

    result = await OfficeTool().execute(
        _create_args("bad"),
        _context(workspace, coordinator, repairer=repairer),
    )

    assert not result.success
    assert result.metadata["office_validation_failure"]["repair_status"] == (
        "rejected"
    )
    assert "RuntimeError" in caplog.text
    assert all(value not in caplog.text for value in private_values)
    serialized_result = json.dumps(result.metadata, ensure_ascii=False, default=str)
    assert all(value not in serialized_result for value in private_values)


@pytest.mark.asyncio
@pytest.mark.parametrize("attack", ["path", "operation", "seal", "image"])
async def test_repair_rejects_identity_authority_and_new_image_reads(
    workspace: Path,
    office_authoring_gate: None,
    attack: str,
) -> None:
    replacement = _create_args("bad")
    if attack == "path":
        replacement["file_path"] = "other.docx"
    elif attack == "operation":
        replacement["operation"] = "edit"
    elif attack == "seal":
        replacement["seal"] = "caller-supplied"  # type: ignore[assignment]
    else:
        image = workspace / "new-image.png"
        Image.new("RGB", (8, 8), color=(10, 20, 30)).save(image)
        replacement["document"] = {
            "title": "bad",
            "images": [{"path": str(image)}],
        }
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _Repairer(
        [replacement],
        preserve_raw_target=attack == "path",
    )

    result = await OfficeTool().execute(
        _create_args("bad"),
        _context(workspace, coordinator, repairer=repairer),
    )

    assert not result.success
    assert len(repairer.requests) == 1
    assert len(coordinator.sessions) == 1
    assert result.metadata["office_validation_failure"]["repair_status"] == (
        "rejected"
    )
    assert not _target(workspace).exists()


@pytest.mark.asyncio
async def test_edit_reset_preserves_concurrent_visible_baseline_change(
    workspace: Path,
    office_authoring_gate: None,
) -> None:
    target = _target(workspace)
    target.parent.mkdir()
    document = Document()
    office_module._drop_default_docx_custom_xml(document)
    document.add_paragraph("baseline")
    document.save(target)
    concurrent = Document()
    office_module._drop_default_docx_custom_xml(concurrent)
    concurrent.add_paragraph("concurrent user edit")
    concurrent_path = target.with_name("concurrent.docx")
    concurrent.save(concurrent_path)

    def replace_visible(_attempt: int, _request: object) -> None:
        os.replace(concurrent_path, target)

    repairer = _Repairer(
        [_edit_layout_args("baseline", "bad", 1.0)],
        before_return=replace_visible,
    )
    coordinator = _Coordinator(["fail", "pass"])

    result = await OfficeTool().execute(
        _edit_layout_args("baseline", "bad", 1.5),
        _context(workspace, coordinator, repairer=repairer),
    )

    assert not result.success
    assert "concurrent user edit" in _document_text(target)
    assert "baseline changed" in (result.error or "").lower()
    assert str(workspace) not in (result.error or "")
    assert len(coordinator.sessions) == 1


@pytest.mark.asyncio
async def test_cancellation_waits_for_repair_worker_before_cleanup(
    workspace: Path,
    office_authoring_gate: None,
) -> None:
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _BlockingRepairer(_create_args("good"))
    task = asyncio.create_task(
        OfficeTool().execute(
            _create_args("bad"),
            _context(workspace, coordinator, repairer=repairer),
        )
    )
    await asyncio.wait_for(repairer.entered.wait(), timeout=5)

    task.cancel()
    await asyncio.sleep(0)
    assert not task.done()
    repairer.release.set()
    with pytest.raises(asyncio.CancelledError):
        await asyncio.wait_for(task, timeout=5)

    assert not _target(workspace).exists()
    private = workspace.parent / "private"
    assert not list(private.rglob("tx-*"))


def test_production_repair_timeout_is_code_owned_and_bounded() -> None:
    assert 1 <= office_module._OFFICE_REPAIR_TIMEOUT_SECONDS <= 300
    assert 0 < office_module._OFFICE_REPAIR_SETTLEMENT_GRACE_SECONDS <= 10


@pytest.mark.asyncio
async def test_never_ending_repairer_times_out_without_changing_original(
    workspace: Path,
    office_authoring_gate: None,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    target = _target(workspace)
    target.parent.mkdir()
    document = Document()
    office_module._drop_default_docx_custom_xml(document)
    document.add_paragraph("baseline")
    document.save(target)
    before = target.read_bytes()
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _CancellationSuppressingRepairer(
        _edit_args("baseline", "repaired")
    )
    caplog.set_level("WARNING", logger=office_module.__name__)
    _install_test_repair_limits(monkeypatch, timeout=0.04, grace=0.03)

    started = asyncio.get_running_loop().time()
    result = await asyncio.wait_for(
        OfficeTool().execute(
            _edit_args("baseline", "bad"),
            _context(workspace, coordinator, repairer=repairer),
        ),
        timeout=1,
    )
    elapsed = asyncio.get_running_loop().time() - started

    assert elapsed < 0.5
    assert not result.success
    assert result.metadata["office_validation_failure"]["repair_status"] == (
        "timeout"
    )
    assert target.read_bytes() == before
    assert len(coordinator.sessions) == 1
    assert repairer.cancelled.is_set()
    request = repairer.requests[0]
    assert not any(
        isinstance(getattr(request, field.name), Path)
        for field in fields(request)
    )
    repairer.release.set()
    await asyncio.wait_for(repairer.finished.wait(), timeout=1)
    await asyncio.sleep(0)
    assert "private repair failure" not in caplog.text
    assert "/staging/secret" not in caplog.text


@pytest.mark.asyncio
async def test_user_cancellation_is_bounded_when_repairer_swallows_cancel(
    workspace: Path,
    office_authoring_gate: None,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    target = _target(workspace)
    target.parent.mkdir()
    document = Document()
    office_module._drop_default_docx_custom_xml(document)
    document.add_paragraph("baseline")
    document.save(target)
    before = target.read_bytes()
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _CancellationSuppressingRepairer(
        _edit_args("baseline", "repaired")
    )
    caplog.set_level("WARNING", logger=office_module.__name__)
    _install_test_repair_limits(monkeypatch, timeout=0.5, grace=0.03)
    task = asyncio.create_task(
        OfficeTool().execute(
            _edit_args("baseline", "bad"),
            _context(workspace, coordinator, repairer=repairer),
        )
    )
    await asyncio.wait_for(repairer.entered.wait(), timeout=1)

    started = asyncio.get_running_loop().time()
    task.cancel()
    with pytest.raises(asyncio.CancelledError):
        await asyncio.wait_for(task, timeout=1)
    elapsed = asyncio.get_running_loop().time() - started

    assert elapsed < 0.5
    assert target.read_bytes() == before
    assert repairer.cancelled.is_set()
    private = workspace.parent / "private"
    assert not list(private.rglob("tx-*"))
    repairer.release.set()
    await asyncio.wait_for(repairer.finished.wait(), timeout=1)
    await asyncio.sleep(0)
    assert "private repair failure" not in caplog.text
    assert "/staging/secret" not in caplog.text


@pytest.mark.asyncio
async def test_slow_repair_within_server_deadline_completes_normally(
    workspace: Path,
    office_authoring_gate: None,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    coordinator = _Coordinator(["fail", "pass"])
    repairer = _SlowRepairer(
        _create_layout_args("fixed content", 1.0),
        delay=0.03,
    )
    _install_test_repair_limits(monkeypatch, timeout=0.2, grace=0.03)

    result = await OfficeTool().execute(
        _create_layout_args("fixed content", 1.5),
        _context(workspace, coordinator, repairer=repairer),
    )

    assert result.success
    assert "fixed content" in _document_text(_target(workspace))
    assert result.metadata["office_validation_repair_attempts"] == 1


def _transaction_context(workspace: Path) -> ToolContext:
    return ToolContext(
        session_id="session",
        message_id="message",
        agent=AgentInfo(name="test", description="", mode="primary"),
        call_id="call",
        workspace=str(workspace),
        root_turn_id="root",
        turn_run_id="run",
        checkpoint_id="checkpoint",
        workspace_instance_id="workspace-instance",
    )


def _raw_seal(
    staged: Path,
    target: Path,
    *,
    validation_generation: str,
) -> OfficeDraftSeal:
    root_info = staged.lstat()
    source_info = target.lstat()
    payload = target.read_bytes()
    return OfficeDraftSeal(
        relative_path=target.relative_to(staged).as_posix(),
        source_sha256=hashlib.sha256(payload).hexdigest(),
        source_mode=stat.S_IMODE(source_info.st_mode),
        source_size=source_info.st_size,
        root_identity=(root_info.st_dev, root_info.st_ino),
        source_identity=(source_info.st_dev, source_info.st_ino),
        validation_generation=validation_generation,
        renderer_id="renderer",
        renderer_version="1",
        font_digest="f" * 64,
        parameters_version="1",
        parameters_sha256="a" * 64,
        quality="authoritative",
        cache_key="b" * 64,
    )


def test_wmt_create_reset_preserves_armed_root_and_deletes_candidate(
    workspace: Path,
) -> None:
    target = workspace / "nested" / "report.docx"
    transaction = WorkspaceMutationTransaction(
        workspace,
        _transaction_context(workspace),
        operation="office.create",
        storage_root=workspace.parent / "private",
    )
    staged = transaction.prepare_paths([target])
    view = transaction.arm_office_precommit_validation(target)
    view.staged_target.parent.mkdir()
    view.staged_target.write_bytes(b"candidate")

    reset_view = transaction.reset_office_precommit_target(target)

    assert not view.staged_target.exists()
    assert view.staged_target.parent.is_dir()
    root_info = staged.lstat()
    assert view.staged_root_identity == (root_info.st_dev, root_info.st_ino)
    assert reset_view.validation_generation != view.validation_generation
    transaction.abort()


def test_wmt_edit_reset_rejects_old_seal_even_for_identical_candidate_bytes(
    workspace: Path,
) -> None:
    target = workspace / "report.docx"
    target.write_bytes(b"baseline")
    transaction = WorkspaceMutationTransaction(
        workspace,
        _transaction_context(workspace),
        operation="office.edit",
        storage_root=workspace.parent / "private",
    )
    staged = transaction.prepare_paths([target])
    view = transaction.arm_office_precommit_validation(target)
    view.staged_target.write_bytes(b"candidate")
    stale_seal = _raw_seal(
        staged,
        view.staged_target,
        validation_generation=view.validation_generation,
    )

    reset_view = transaction.reset_office_precommit_target(target)
    assert view.staged_target.read_bytes() == b"baseline"
    view.staged_target.write_bytes(b"candidate")
    # Make every filesystem/content field in the stale seal match the new
    # candidate explicitly.  The transaction generation alone must still
    # reject it, even on filesystems that immediately reuse the old inode.
    source_info = view.staged_target.lstat()
    stale_seal = replace(
        stale_seal,
        source_identity=(source_info.st_dev, source_info.st_ino),
    )
    assert reset_view.validation_generation != view.validation_generation

    with pytest.raises(WorkspacePrecommitSealError, match="generation is stale"):
        transaction.commit_with_precommit_office_seal(stale_seal)
    assert target.read_bytes() == b"baseline"


def test_wmt_edit_reset_accepts_only_a_new_generation_seal(
    workspace: Path,
) -> None:
    target = workspace / "report.docx"
    target.write_bytes(b"baseline")
    transaction = WorkspaceMutationTransaction(
        workspace,
        _transaction_context(workspace),
        operation="office.edit",
        storage_root=workspace.parent / "private",
    )
    staged = transaction.prepare_paths([target])
    view = transaction.arm_office_precommit_validation(target)
    view.staged_target.write_bytes(b"first candidate")

    reset_view = transaction.reset_office_precommit_target(target)
    reset_view.staged_target.write_bytes(b"second candidate")
    fresh_seal = _raw_seal(
        staged,
        reset_view.staged_target,
        validation_generation=reset_view.validation_generation,
    )

    transaction.commit_with_precommit_office_seal(fresh_seal)

    assert target.read_bytes() == b"second candidate"


def test_wmt_failed_edit_reset_retires_old_generation(
    workspace: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    target = workspace / "report.docx"
    target.write_bytes(b"baseline")
    transaction = WorkspaceMutationTransaction(
        workspace,
        _transaction_context(workspace),
        operation="office.edit",
        storage_root=workspace.parent / "private",
    )
    staged = transaction.prepare_paths([target])
    view = transaction.arm_office_precommit_validation(target)
    view.staged_target.write_bytes(b"candidate")
    stale_seal = _raw_seal(
        staged,
        view.staged_target,
        validation_generation=view.validation_generation,
    )

    def fail_after_candidate_retirement(
        workspace_fd: int,
        relative: str,
        destination: Path,
    ) -> tuple[object, int]:
        del workspace_fd, relative, destination
        raise WorkspacePrecommitSealError("injected reset copy failure")

    monkeypatch.setattr(
        "app.tool.workspace_transaction._copy_regular_file_at_relative",
        fail_after_candidate_retirement,
    )
    with pytest.raises(WorkspacePrecommitSealError, match="injected"):
        transaction.reset_office_precommit_target(target)

    view.staged_target.write_bytes(b"candidate")
    source_info = view.staged_target.lstat()
    stale_seal = replace(
        stale_seal,
        source_identity=(source_info.st_dev, source_info.st_ino),
    )
    with pytest.raises(WorkspacePrecommitSealError, match="generation is retired"):
        transaction.commit_with_precommit_office_seal(stale_seal)
    assert target.read_bytes() == b"baseline"


@pytest.mark.parametrize("mutation", ["sha", "identity"])
def test_wmt_edit_reset_requires_visible_baseline_sha_and_identity(
    workspace: Path,
    mutation: str,
) -> None:
    target = workspace / "report.docx"
    target.write_bytes(b"baseline")
    transaction = WorkspaceMutationTransaction(
        workspace,
        _transaction_context(workspace),
        operation="office.edit",
        storage_root=workspace.parent / "private",
    )
    transaction.prepare_paths([target])
    view = transaction.arm_office_precommit_validation(target)
    view.staged_target.write_bytes(b"candidate")
    if mutation == "sha":
        target.write_bytes(b"changed")
    else:
        replacement = target.with_name("replacement.docx")
        replacement.write_bytes(b"baseline")
        os.replace(replacement, target)

    with pytest.raises(WorkspacePrecommitSealError, match="baseline changed"):
        transaction.reset_office_precommit_target(target)

    assert view.staged_target.read_bytes() == b"candidate"
    transaction.abort()
