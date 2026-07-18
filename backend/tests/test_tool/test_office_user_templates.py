"""Targeted OfficeTool coverage for approved workspace user templates."""

from __future__ import annotations

import io
import json
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from sqlalchemy import select

from app import release_features
from app.models.office_user_template import OfficeUserTemplate
from app.models.session import Session
from app.office_rendering import OfficeRenderCache, RendererDescriptor
from app.office_templates.user import (
    UserOfficeTemplateService,
    UserTemplatePlaceholder,
)
from app.office_validation import (
    OfficeDraftValidationResult,
    OfficeValidationReport,
    ValidationCheck,
)
from app.office_validation.draft import OfficeDraftValidationService
from app.office_validation.precommit import (
    DeterministicOfficePrecommitCoordinator,
)
from app.schemas.agent import AgentInfo
from app.storage.checkpoints import register_workspace_instance
from app.tool.builtin.office import OfficeTool
from app.tool.context import ToolContext
from app.tool.workspace_transaction import (
    committed_checkpoint_journal_action,
    committed_checkpoint_journal_metadata,
    list_committed_checkpoint_journals,
)
from tests.test_office_rendering.helpers import FakeProvider
from tests.test_office_templates.helpers import make_docx_template


pytestmark = pytest.mark.asyncio

_PLACEHOLDERS = ("body", "client", "footer", "header", "table")
_VALID_VALUES = {
    "body": "OK",
    "client": "Acme",
    "footer": "Internal",
    "header": "Quarterly",
    "table": "Ready",
}


def _release_user_templates(monkeypatch: pytest.MonkeyPatch) -> None:
    for name in (
        "V11_CHECKPOINTS_RELEASED",
        "V11_REWIND_RELEASED",
        "V11_VALIDATION_AGENT_RELEASED",
        "V11_OFFICE_V2_RELEASED",
        "V11_USER_OFFICE_TEMPLATES_BETA_RELEASED",
    ):
        monkeypatch.setattr(release_features, name, True)


def _placeholder_schema() -> tuple[UserTemplatePlaceholder, ...]:
    return tuple(
        UserTemplatePlaceholder(
            name=name,
            value_type="text",
            required=True,
            min_chars=2 if name == "body" else 1,
            max_chars=4 if name == "body" else 40,
            description=f"Value for {name}",
        )
        for name in _PLACEHOLDERS
    )


class _NeverFallbackPolicies:
    """The user-template trusted plan must bypass ordinary policy lookup."""

    def resolve_create(self, request: object) -> object:
        raise AssertionError("user template unexpectedly used fallback create policy")

    def resolve_edit(self, request: object, baseline: object) -> object:
        raise AssertionError("user template unexpectedly used fallback edit policy")


@dataclass(slots=True)
class _TemplateEnvironment:
    workspace: Path
    private: Path
    session_id: str
    workspace_instance_id: str
    template_ref: str
    revision: int
    state_version: int
    source_sha256: str
    manifest_sha256: str
    service: UserOfficeTemplateService
    draft: OfficeDraftValidationService
    coordinator: Any
    session_factory: Any

    def context(
        self,
        *,
        suffix: str = "use",
        workspace: Path | None = None,
        session_id: str | None = None,
        workspace_instance_id: str | None = None,
        coordinator: Any | None = None,
        repairer: Any | None = None,
    ) -> ToolContext:
        ctx = ToolContext(
            session_id=session_id or self.session_id,
            message_id=f"user-template-message-{suffix}",
            agent=AgentInfo(name="test", description="", mode="primary"),
            call_id=f"user-template-call-{suffix}",
            language="en",
            workspace=str((workspace or self.workspace).resolve()),
            root_turn_id=f"user-template-root-{suffix}",
            turn_run_id=f"user-template-run-{suffix}",
            checkpoint_id=f"user-template-checkpoint-{suffix}",
            workspace_instance_id=(
                workspace_instance_id or self.workspace_instance_id
            ),
        )
        state: dict[str, object] = {
            "session_factory": self.session_factory,
            "office_user_template_service": self.service,
            "office_precommit_coordinator": coordinator or self.coordinator,
        }
        if repairer is not None:
            state["office_precommit_repairer"] = repairer
        ctx._app_state = state  # type: ignore[attr-defined]
        return ctx

    def args(
        self,
        filename: str,
        *,
        values: dict[str, str] | None = None,
        state_version: int | None = None,
    ) -> dict[str, object]:
        return {
            "file_path": filename,
            "operation": "create",
            "user_template": {
                "template_ref": self.template_ref,
                "revision": self.revision,
                "expected_state_version": (
                    self.state_version if state_version is None else state_version
                ),
                "values": dict(_VALID_VALUES if values is None else values),
            },
        }

    def output(self, filename: str) -> Path:
        return self.workspace / "suxiaoyou_written" / filename


async def _create_workspace(
    session_factory: Any,
    tmp_path: Path,
    *,
    suffix: str,
) -> tuple[str, str, Path]:
    workspace = tmp_path / f"workspace-{suffix}"
    workspace.mkdir()
    session_id = f"user-template-session-{suffix}"
    async with session_factory() as db:
        async with db.begin():
            db.add(
                Session(
                    id=session_id,
                    directory=str(workspace.resolve()),
                    title="Office user-template tool test",
                )
            )
            await db.flush()
            instance = await register_workspace_instance(
                db,
                workspace,
                kind="direct",
                created_by_session_id=session_id,
            )
            instance_id = instance.id
    return session_id, instance_id, workspace


async def _approved_environment(
    session_factory: Any,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    *,
    suffix: str,
) -> _TemplateEnvironment:
    _release_user_templates(monkeypatch)
    private = tmp_path / f"private-{suffix}"
    private.mkdir()
    monkeypatch.setenv("SUXIAOYOU_PRIVATE_DATA_DIR", str(private))
    session_id, instance_id, workspace = await _create_workspace(
        session_factory,
        tmp_path,
        suffix=suffix,
    )

    descriptor = RendererDescriptor(
        renderer_id=f"user-template-renderer-{suffix}",
        renderer_version="1.0.0",
        font_digest="a" * 64,
        quality="authoritative",
    )
    provider = FakeProvider(descriptor)
    draft = OfficeDraftValidationService(
        cache=OfficeRenderCache((tmp_path / f"render-cache-{suffix}").absolute()),
        provider=provider,
        parameters_version="user-template-test-v1",
        parameters={"dpi": 144},
    )
    service = UserOfficeTemplateService(
        (tmp_path / f"template-registry-{suffix}").absolute(),
        draft_validation=draft,
    )
    candidate = await service.validate_and_register(
        io.BytesIO(make_docx_template()),
        filename="approved.docx",
        display_name="Approved quarterly template",
        placeholders=_placeholder_schema(),
    )
    assert candidate.status == "needs_confirmation"
    now = datetime.now(timezone.utc)
    async with session_factory() as db:
        async with db.begin():
            db.add(
                OfficeUserTemplate(
                    template_ref=candidate.template_ref,
                    revision=candidate.revision,
                    state_version=2,
                    workspace_instance_id=instance_id,
                    created_by_session_id=session_id,
                    import_idempotency_key=f"import-{suffix}",
                    import_request_sha256=candidate.import_request_sha256,
                    display_name=candidate.display_name,
                    format=candidate.format,
                    source_sha256=candidate.source_sha256,
                    source_size_bytes=candidate.source_size_bytes,
                    manifest_sha256=candidate.manifest_sha256,
                    placeholder_schema=[
                        field.to_dict() for field in candidate.placeholder_schema
                    ],
                    placeholder_parts=list(candidate.placeholder_parts),
                    allowed_operations=list(candidate.allowed_operations),
                    status="approved",
                    render_quality=candidate.render_manifest.quality,
                    renderer_id=candidate.render_manifest.renderer_id,
                    renderer_version=candidate.render_manifest.renderer_version,
                    font_digest=candidate.render_manifest.font_digest,
                    render_parameters_version=(
                        candidate.render_manifest.parameters_version
                    ),
                    render_parameters_sha256=(
                        candidate.render_manifest.parameters_sha256
                    ),
                    render_cache_key=candidate.render_manifest.cache_key,
                    render_manifest_sha256=candidate.render_manifest_sha256,
                    render_page_count=len(candidate.render_manifest.pages),
                    validation_report=dict(candidate.validation_report),
                    time_approved=now,
                )
            )
    coordinator = DeterministicOfficePrecommitCoordinator(
        service=draft,
        policies=_NeverFallbackPolicies(),
    )
    return _TemplateEnvironment(
        workspace=workspace,
        private=private,
        session_id=session_id,
        workspace_instance_id=instance_id,
        template_ref=candidate.template_ref,
        revision=candidate.revision,
        state_version=2,
        source_sha256=candidate.source_sha256,
        manifest_sha256=candidate.manifest_sha256,
        service=service,
        draft=draft,
        coordinator=coordinator,
        session_factory=session_factory,
    )


async def test_user_template_schema_requires_its_composed_beta_gate(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    for name in (
        "V11_CHECKPOINTS_RELEASED",
        "V11_REWIND_RELEASED",
        "V11_VALIDATION_AGENT_RELEASED",
        "V11_OFFICE_V2_RELEASED",
    ):
        monkeypatch.setattr(release_features, name, True)
    monkeypatch.setattr(
        release_features,
        "V11_USER_OFFICE_TEMPLATES_BETA_RELEASED",
        False,
    )
    closed = OfficeTool().parameters_schema()
    assert "first_party_template" in closed["properties"]
    assert "user_template" not in closed["properties"]

    monkeypatch.setattr(
        release_features,
        "V11_USER_OFFICE_TEMPLATES_BETA_RELEASED",
        True,
    )
    released = OfficeTool().parameters_schema()
    contract = released["properties"]["user_template"]
    assert contract["additionalProperties"] is False
    assert contract["required"] == [
        "template_ref",
        "revision",
        "expected_state_version",
        "values",
    ]
    assert contract["properties"]["template_ref"]["pattern"] == (
        r"^utpl-[0-9a-z]{26}$"
    )

    monkeypatch.setattr(release_features, "V11_REWIND_RELEASED", False)
    dependency_closed = OfficeTool().parameters_schema()
    assert "user_template" not in dependency_closed["properties"]


async def test_user_template_is_workspace_scoped_and_state_cas_bound(
    session_factory: Any,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    env = await _approved_environment(
        session_factory,
        tmp_path,
        monkeypatch,
        suffix="scope",
    )
    other_session, other_instance, other_workspace = await _create_workspace(
        session_factory,
        tmp_path,
        suffix="other",
    )
    tool = OfficeTool()

    cross_workspace = await tool.execute(
        env.args("cross-workspace.docx"),
        env.context(
            suffix="cross",
            workspace=other_workspace,
            session_id=other_session,
            workspace_instance_id=other_instance,
        ),
    )
    assert not cross_workspace.success
    assert not (
        other_workspace / "suxiaoyou_written" / "cross-workspace.docx"
    ).exists()

    stale = await tool.execute(
        env.args("stale-state.docx", state_version=1),
        env.context(suffix="stale"),
    )
    assert not stale.success
    assert not env.output("stale-state.docx").exists()


@pytest.mark.parametrize("body", ["", "12345"])
async def test_user_template_enforces_each_approved_field_bound(
    body: str,
    session_factory: Any,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    env = await _approved_environment(
        session_factory,
        tmp_path,
        monkeypatch,
        suffix=f"bounds-{len(body)}",
    )
    values = {**_VALID_VALUES, "body": body}
    filename = f"invalid-{len(body)}.docx"
    result = await OfficeTool().execute(
        env.args(filename, values=values),
        env.context(suffix=f"bounds-{len(body)}"),
    )

    assert not result.success
    assert "text bounds" in (result.error or "")
    assert not env.output(filename).exists()


async def test_user_template_success_commits_path_free_checkpoint_evidence(
    session_factory: Any,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    env = await _approved_environment(
        session_factory,
        tmp_path,
        monkeypatch,
        suffix="success",
    )
    result = await OfficeTool().execute(
        env.args("quarterly.docx"),
        env.context(suffix="success"),
    )

    assert result.success, result.error
    output = env.output("quarterly.docx")
    assert output.is_file()
    document = Document(str(output))
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Acme" in body_text
    assert "OK" in body_text
    assert "Ready" in document.tables[0].cell(0, 0).text
    assert "Quarterly" in document.sections[0].header.paragraphs[0].text
    assert "Internal" in document.sections[0].footer.paragraphs[0].text
    assert "{{" not in output.read_bytes().decode("latin-1")

    metadata = result.metadata
    assert metadata["file_path"] == "suxiaoyou_written/quarterly.docx"
    assert metadata["written_files"] == ["suxiaoyou_written/quarterly.docx"]
    assert metadata["user_template_ref"] == env.template_ref
    assert metadata["user_template_revision"] == 1
    assert metadata["user_template_state_version"] == 2
    assert metadata["user_template_source_sha256"] == env.source_sha256
    assert metadata["user_template_manifest_sha256"] == env.manifest_sha256
    assert metadata["office_validation_checkpoint_id"] == (
        "user-template-checkpoint-success"
    )
    assert metadata["office_validation_root_turn_id"] == (
        "user-template-root-success"
    )
    report = metadata["_office_validation_report"]
    assert report["verdict"] == "pass"
    assert report["baseline_sha256"] == env.source_sha256
    assert str(tmp_path) not in json.dumps(metadata, ensure_ascii=False)

    mutations = metadata["workspace_mutations"]
    file_mutation = next(
        item
        for item in mutations
        if item["relative_path"] == "suxiaoyou_written/quarterly.docx"
    )
    assert file_mutation["operation"] == "created"
    token = metadata["_checkpoint_journal"]
    journals = list_committed_checkpoint_journals(storage_root=env.private)
    assert [item[0] for item in journals] == [token]
    runtime, durable = committed_checkpoint_journal_metadata(journals[0][1])
    assert runtime["checkpoint_id"] == "user-template-checkpoint-success"
    assert runtime["root_turn_id"] == "user-template-root-success"
    assert runtime["workspace_instance_id"] == env.workspace_instance_id
    assert durable["workspace_mutations"] == mutations
    assert committed_checkpoint_journal_action(journals[0][1]) == (
        "turn_commit",
        (),
    )


class _AfterValidationSession:
    def __init__(self, inner: Any, callback: Any) -> None:
        self._inner = inner
        self._callback = callback

    async def validate_candidate(self) -> OfficeDraftValidationResult:
        result = await self._inner.validate_candidate()
        await self._callback()
        return result

    def consume_commit_seal(self, result: OfficeDraftValidationResult) -> Any:
        return self._inner.consume_commit_seal(result)

    def mark_committed(self, result: OfficeDraftValidationResult) -> None:
        self._inner.mark_committed(result)

    def abort(self) -> None:
        self._inner.abort()


class _AfterValidationCoordinator:
    def __init__(self, inner: Any, callback: Any) -> None:
        self._inner = inner
        self._callback = callback

    async def begin(self, *, request: Any, view: Any) -> Any:
        inner = await self._inner.begin(request=request, view=view)
        return _AfterValidationSession(inner, self._callback)


async def test_user_template_tombstone_after_validation_never_publishes(
    session_factory: Any,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    env = await _approved_environment(
        session_factory,
        tmp_path,
        monkeypatch,
        suffix="tombstone-race",
    )

    async def tombstone() -> None:
        async with session_factory() as db:
            async with db.begin():
                record = (
                    await db.execute(
                        select(OfficeUserTemplate).where(
                            OfficeUserTemplate.template_ref == env.template_ref,
                            OfficeUserTemplate.revision == env.revision,
                        )
                    )
                ).scalar_one()
                record.status = "tombstoned"
                record.state_version += 1
                record.time_tombstoned = datetime.now(timezone.utc)

    coordinator = _AfterValidationCoordinator(env.coordinator, tombstone)
    result = await OfficeTool().execute(
        env.args("raced.docx"),
        env.context(suffix="race", coordinator=coordinator),
    )

    assert not result.success
    assert not env.output("raced.docx").exists()
    assert list_committed_checkpoint_journals(storage_root=env.private) == []
    async with session_factory() as db:
        record = (
            await db.execute(
                select(OfficeUserTemplate).where(
                    OfficeUserTemplate.template_ref == env.template_ref,
                    OfficeUserTemplate.revision == env.revision,
                )
            )
        ).scalar_one()
        assert record.status == "tombstoned"
        assert record.state_version == 3


class _FailingValidationSession:
    def __init__(self, inner: Any) -> None:
        self._inner = inner

    async def validate_candidate(self) -> OfficeDraftValidationResult:
        passing = await self._inner.validate_candidate()
        check = ValidationCheck(
            code="forced_failure",
            outcome="fail",
            message="Test-only authoritative rejection.",
        )
        report = OfficeValidationReport(
            document_format=passing.report.document_format,
            baseline_sha256=passing.report.baseline_sha256,
            candidate_sha256=passing.report.candidate_sha256,
            renderer_id=passing.report.renderer_id,
            renderer_version=passing.report.renderer_version,
            font_digest=passing.report.font_digest,
            verdict="fail",
            checks=(check,),
            checkpoint_id=passing.report.checkpoint_id,
            root_turn_id=passing.report.root_turn_id,
        )
        return OfficeDraftValidationResult(report=report, candidate=passing.candidate)

    def consume_commit_seal(self, result: OfficeDraftValidationResult) -> Any:
        raise AssertionError("a failed user-template candidate cannot be consumed")

    def mark_committed(self, result: OfficeDraftValidationResult) -> None:
        raise AssertionError("a failed user-template candidate cannot be committed")

    def abort(self) -> None:
        self._inner.abort()


class _FailingValidationCoordinator:
    def __init__(self, inner: Any) -> None:
        self._inner = inner

    async def begin(self, *, request: Any, view: Any) -> Any:
        inner = await self._inner.begin(request=request, view=view)
        return _FailingValidationSession(inner)


class _RepairSpy:
    def __init__(self) -> None:
        self.calls = 0

    async def repair(self, request: Any) -> dict[str, object]:
        self.calls += 1
        raise AssertionError("user-template validation must fail closed")


async def test_user_template_validation_failure_does_not_call_repairer(
    session_factory: Any,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    env = await _approved_environment(
        session_factory,
        tmp_path,
        monkeypatch,
        suffix="repair-closed",
    )
    repairer = _RepairSpy()
    result = await OfficeTool().execute(
        env.args("rejected.docx"),
        env.context(
            suffix="repair",
            coordinator=_FailingValidationCoordinator(env.coordinator),
            repairer=repairer,
        ),
    )

    assert not result.success
    assert repairer.calls == 0
    assert not env.output("rejected.docx").exists()
    failure = result.metadata["office_validation_failure"]
    assert failure["repair_attempts"] == 0
    assert failure["repair_status"] == "unavailable"
    assert list_committed_checkpoint_journals(storage_root=env.private) == []
